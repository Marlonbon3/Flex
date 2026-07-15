from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ──────────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x0A, 0x1B, 0x5B)
AZUL_MEDIO   = RGBColor(0x0A, 0x46, 0xFF)
GRIS_TEXTO   = RGBColor(0x33, 0x33, 0x33)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
VERDE        = RGBColor(0x10, 0xB9, 0x81)
AMARILLO     = RGBColor(0xF5, 0x9E, 0x0B)
ROJO         = RGBColor(0xDC, 0x26, 0x26)
GRIS_CLARO   = RGBColor(0xF0, 0xF4, 0xFF)

doc = Document()

# ── Márgenes ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = AZUL_OSCURO
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = AZUL_MEDIO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS_TEXTO
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(11)
    run.font.color.rgb = color or GRIS_TEXTO
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = GRIS_TEXTO
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GRIS_TEXTO
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRIS_TEXTO
    p.paragraph_format.left_indent  = Cm(0.63 + level * 0.63)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = GRIS_TEXTO
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GRIS_TEXTO
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRIS_TEXTO
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(text, color_hex='0A46FF'):
    p = doc.add_paragraph()
    run = p.add_run('  ' + text + '  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_role_table(roles):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Rol', 'Descripción', 'Permisos clave']):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(10)
        set_cell_bg(hdr[i], '0A1B5B')
    for rol, desc, perms in roles:
        row = tbl.add_row().cells
        row[0].text = rol
        row[1].text = desc
        row[2].text = perms
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FLEXTRONICS TECHNOLOGIES')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZUL_OSCURO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('San Luis Río Colorado, Sonora, México')
r2.font.size = Pt(13); r2.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('MANUAL DE USUARIO')
r3.bold = True; r3.font.size = Pt(32); r3.font.color.rgb = AZUL_MEDIO

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Sistema de Gestión de Recibos y Contenedores')
r4.font.size = Pt(16); r4.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('Flex WebApp  ·  v1.0  ·  2026')
r5.font.size = Pt(11); r5.font.color.rgb = RGBColor(0x77,0x77,0x77)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════
add_h1('Índice')
secciones = [
    ('1', 'Descripción general del sistema'),
    ('2', 'Acceso al sistema (Autenticación)'),
    ('3', 'Roles y permisos'),
    ('4', 'Navegación — Barra lateral'),
    ('5', 'Pantalla de Inicio'),
    ('6', 'Módulo: Contenedores'),
    ('  6.1', 'Ver la lista de contenedores activos'),
    ('  6.2', 'Agregar un nuevo contenedor — Paso 1: Información Básica'),
    ('  6.3', 'Agregar un nuevo contenedor — Paso 2: Inspección de Trailer'),
    ('  6.4', 'Agregar un nuevo contenedor — Paso 3: Documentos'),
    ('  6.5', 'Editar un contenedor existente'),
    ('  6.6', 'Archivar un contenedor'),
    ('  6.7', 'Vaciar campos opcionales'),
    ('  6.8', 'Eliminar un contenedor'),
    ('7', 'Módulo: Archivo'),
    ('  7.1', 'Filtros disponibles'),
    ('  7.2', 'Generar y descargar PDF'),
    ('  7.3', 'Editar registro archivado'),
    ('  7.4', 'Eliminar del archivo'),
    ('  7.5', 'Exportar a Excel'),
    ('8', 'Módulo: Reportes'),
    ('  8.1', 'KPIs — Tarjetas de indicadores'),
    ('  8.2', 'Panel expandible de detalle'),
    ('  8.3', 'Progreso de completado (gráfica de anillo)'),
    ('  8.4', 'Contenedores por mes'),
    ('  8.5', 'Gráficas de barras (Turno / Origen / Empresa)'),
    ('9', 'Módulo: Entrega de Turno'),
    ('  9.1', 'Visualizar el turno activo'),
    ('  9.2', 'Guardar entrega de turno'),
    ('  9.3', 'Exportar entrega a Excel'),
    ('  9.4', 'Eliminar entrega guardada'),
    ('10', 'Panel de Administración (solo Administradores)'),
    ('  10.1', 'Configuración de Listas'),
    ('  10.2', 'Gestión de Usuarios'),
    ('11', 'Preguntas frecuentes'),
]
for num, titulo in secciones:
    p = doc.add_paragraph()
    r1 = p.add_run(num.ljust(6))
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AZUL_MEDIO
    r2 = p.add_run(titulo)
    r2.font.size = Pt(10); r2.font.color.rgb = GRIS_TEXTO
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 1 — DESCRIPCIÓN GENERAL
# ═══════════════════════════════════════════════════════════════════════════
add_h1('1. Descripción general del sistema')
add_body(
    'Flex WebApp es un sistema digital de gestión de recibos y contenedores desarrollado para '
    'Flextronics Technologies, sede San Luis Río Colorado, Sonora. Permite registrar, inspeccionar, '
    'archivar y reportar el flujo operativo de tráilers y contenedores marítimos que ingresan '
    'al almacén.')

add_h3('Funcionalidades principales')
for f in [
    'Registro de contenedores en tres pasos con validaciones automáticas.',
    'Firma digital del responsable de descarga directamente en pantalla táctil o tableta.',
    'Adjuntar fotos y documentos (imágenes / PDF) a cada contenedor.',
    'Archivo centralizado de contenedores completados con exportación a PDF y Excel.',
    'Dashboard de reportes con KPIs, gráficas de barras y gráfica de anillo.',
    'Módulo de entrega de turno con exportación a Excel.',
    'Panel de administración para gestionar usuarios y listas configurables del formulario.',
    'Control de acceso basado en roles: Administrador, Supervisor y Operador.',
]:
    add_bullet(f)

add_h3('Tecnología')
for t in [
    'Frontend: React 18 + Vite — interfaz responsiva para escritorio y tabletas.',
    'Backend: Node.js + Express — API REST.',
    'Base de datos: Microsoft SQL Server.',
    'Documentos: exportación a PDF (jsPDF) y Excel (ExcelJS).',
]:
    add_bullet(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 2 — AUTENTICACIÓN
# ═══════════════════════════════════════════════════════════════════════════
add_h1('2. Acceso al sistema (Autenticación)')

add_h2('2.1  Iniciar sesión')
add_body('Al abrir la aplicación, el usuario verá la pantalla de inicio de sesión. Solo los usuarios '
         'registrados en el sistema por un Administrador pueden acceder.')

add_numbered('Abrir el navegador e ingresar la dirección del sistema (URL provista por el área de TI).')
add_numbered('Escribir el ', bold_prefix=None)
p = doc.add_paragraph(style='List Number')
r1 = p.add_run('Correo electrónico'); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=GRIS_TEXTO
r2 = p.add_run(' registrado.'); r2.font.size=Pt(11); r2.font.color.rgb=GRIS_TEXTO
p.paragraph_format.left_indent = Cm(0.63); p.paragraph_format.space_after=Pt(2)
p = doc.add_paragraph(style='List Number')
r1 = p.add_run('Contraseña'); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=GRIS_TEXTO
r2 = p.add_run(' asignada.'); r2.font.size=Pt(11); r2.font.color.rgb=GRIS_TEXTO
p.paragraph_format.left_indent = Cm(0.63); p.paragraph_format.space_after=Pt(2)
add_numbered('Presionar el botón "Iniciar sesión".')
add_numbered('Si los datos son correctos, el sistema redirige a la pantalla de Inicio.')

add_note('⚠  Si el usuario está marcado como Inactivo por el Administrador, el acceso '
         'será denegado aunque las credenciales sean correctas.', 'DC2626')

add_h2('2.2  Cerrar sesión')
add_body('Para cerrar sesión de forma segura:')
add_numbered('En la barra lateral izquierda, localizar el botón "Cerrar sesión" en la parte inferior.')
add_numbered('Hacer clic en el botón.')
add_numbered('El sistema redirige de inmediato a la pantalla de inicio de sesión y borra la sesión activa.')

add_note('Siempre cerrar sesión al dejar una estación de trabajo desatendida para proteger la información.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 3 — ROLES Y PERMISOS
# ═══════════════════════════════════════════════════════════════════════════
add_h1('3. Roles y permisos')
add_body('El sistema cuenta con tres roles de usuario. Cada rol define qué acciones puede realizar '
         'y a qué secciones puede acceder el usuario.')

roles = [
    ('🛡️ Administrador',
     'Acceso total al sistema. Puede gestionar usuarios, configurar listas, visualizar reportes y '
     'modificar o eliminar cualquier registro.',
     '✔ Todas las secciones\n✔ Crear / activar / desactivar usuarios\n✔ Configurar listas del formulario\n'
     '✔ Agregar, editar, archivar y eliminar contenedores\n✔ Panel de Administración visible'),

    ('👁️ Supervisor',
     'Puede ver registros, generar reportes, archivar contenedores y supervisar el flujo operativo. '
     'No puede gestionar usuarios ni modificar la configuración del sistema.',
     '✔ Ver contenedores (solo lectura)\n✔ Ver Archivo y generar PDF\n✔ Ver Reportes\n'
     '✔ Ver Entrega de turno y exportar Excel\n✘ No puede agregar / editar / eliminar contenedores\n'
     '✘ No puede guardar entrega de turno ni eliminarla\n✘ Sin acceso al Panel de Administración'),

    ('⚙️ Operador',
     'Registra nuevos contenedores, completa los pasos del proceso y actualiza documentos. '
     'Acceso limitado a las funciones básicas de operación.',
     '✔ Agregar, editar y archivar contenedores\n✔ Subir documentos (Paso 3)\n'
     '✔ Guardar entrega de turno\n✔ Ver Archivo y Reportes\n'
     '✘ Sin acceso al Panel de Administración'),
]

add_role_table(roles)

add_body('')
add_note('El rol se asigna al crear el usuario desde el Panel de Administración y puede verse '
         'reflejado en la barra lateral como una etiqueta de color junto al nombre del usuario.')

add_h2('3.1  Resumen de acceso por módulo')

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Módulo / Acción', 'Administrador', 'Supervisor', 'Operador']
for i, h in enumerate(hdrs):
    tbl2.rows[0].cells[i].text = h
    for run in tbl2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = BLANCO
    set_cell_bg(tbl2.rows[0].cells[i], '0A1B5B')

permisos = [
    ('Inicio (pantalla principal)',    '✔', '✔', '✔'),
    ('Contenedores — Ver lista',       '✔', '✔', '✔'),
    ('Contenedores — Agregar',         '✔', '✘', '✔'),
    ('Contenedores — Editar',          '✔', '✘', '✔'),
    ('Contenedores — Archivar',        '✔', '✘', '✔'),
    ('Contenedores — Vaciar',          '✔', '✘', '✔'),
    ('Contenedores — Eliminar',        '✔', '✘', '✔'),
    ('Archivo — Ver y PDF',            '✔', '✔', '✔'),
    ('Archivo — Editar',               '✔', '✘', '✔'),
    ('Archivo — Eliminar',             '✔', '✘', '✔'),
    ('Archivo — Exportar Excel',       '✔', '✔', '✔'),
    ('Reportes — Ver dashboard',       '✔', '✔', '✔'),
    ('Entrega Turno — Ver',            '✔', '✔', '✔'),
    ('Entrega Turno — Guardar',        '✔', '✘', '✔'),
    ('Entrega Turno — Exportar Excel', '✔', '✔', '✔'),
    ('Entrega Turno — Eliminar',       '✔', '✘', '✔'),
    ('Administración — Panel',         '✔', '✘', '✘'),
    ('Administración — Listas',        '✔', '✘', '✘'),
    ('Administración — Usuarios',      '✔', '✘', '✘'),
]
for accion, adm, sup, op in permisos:
    row = tbl2.add_row().cells
    row[0].text = accion
    row[1].text = adm; row[2].text = sup; row[3].text = op
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            if run.text == '✔': run.font.color.rgb = VERDE
            elif run.text == '✘': run.font.color.rgb = ROJO

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 4 — NAVEGACIÓN
# ═══════════════════════════════════════════════════════════════════════════
add_h1('4. Navegación — Barra lateral')
add_body('Una vez dentro del sistema, la barra lateral izquierda es el punto de navegación principal. '
         'Contiene los accesos a todos los módulos disponibles según el rol del usuario.')

elementos = [
    ('Inicio',                 'Pantalla de bienvenida con carrusel, misión, visión e información de la sede.'),
    ('Contenedores',           'Gestión del flujo activo de tráilers: registrar, editar, archivar y eliminar.'),
    ('Archivo',                'Consulta de contenedores completados y archivados. Exportación a PDF y Excel.'),
    ('Reportes',               'Dashboard con KPIs, gráficas y estadísticas operativas.'),
    ('Entrega de turno',       'Resumen de contenedores completados por turno. Guardado y exportación.'),
    ('Administración',         '(Solo Administradores) Gestión de usuarios y listas configurables.'),
]
for nombre, desc in elementos:
    add_bullet(desc, bold_prefix=f'{nombre}: ')

add_h3('Información del usuario en la barra lateral')
add_body('En la parte inferior de la barra lateral se muestra:')
add_bullet('Avatar con las iniciales del nombre del usuario.')
add_bullet('Nombre completo del usuario.')
add_bullet('Etiqueta de color indicando el rol (Admin / Supervisor / Operador).')
add_bullet('Botón "Cerrar sesión".')

add_h3('Menú hamburguesa (dispositivos móviles y tabletas)')
add_body('En pantallas pequeñas, la barra lateral se oculta y se accede presionando el '
         'ícono de tres líneas (☰) en la esquina superior izquierda. Al abrir el menú, '
         'un fondo oscuro cubre el contenido; tocarlo cierra el menú.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 5 — INICIO
# ═══════════════════════════════════════════════════════════════════════════
add_h1('5. Pantalla de Inicio')
add_body('La pantalla de Inicio es la primera vista al ingresar al sistema. Presenta información '
         'corporativa y contextual de la empresa.')

add_h3('Contenido de la pantalla')
add_bullet('Carrusel de imágenes: tres diapositivas automáticas con mensajes sobre el sistema '
           '(avanza cada 5 segundos, también puede navegarse manualmente con las flechas).')
add_bullet('Misión: descripción de la misión de Flextronics Technologies.')
add_bullet('Visión: descripción de la visión de la empresa.')
add_bullet('Tarjetas informativas de la sede: Empresa, Ubicación, Operaciones, Especialidad, '
           'Tecnología y Presencia Regional.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 6 — CONTENEDORES
# ═══════════════════════════════════════════════════════════════════════════
add_h1('6. Módulo: Contenedores')
add_body('Este módulo es el núcleo operativo del sistema. Aquí se registran y gestionan todos '
         'los tráilers y contenedores que llegan al almacén.')

# ── 6.1 ─────────────────────────────────────────────────────────────────
add_h2('6.1  Ver la lista de contenedores activos')
add_body('Al entrar al módulo se muestra una tabla con todos los contenedores activos (no archivados). '
         'Las columnas disponibles son:')

cols = [
    ('TRAILER NO.',     'Número identificador del tráiler.'),
    ('TIPO DE TRAILER', 'Tipo de tráiler registrado.'),
    ('CONTENEDOR',      'Tipo de contenedor marítimo (Sea Container Type).'),
    ('PUERTO DE ENTRADA','Puerto de entrada declarado.'),
    ('LLEGADA',         'Fecha y hora en que se creó el registro.'),
    ('STATUS',          'Estado actual: "En proceso", "Completado" o "EMPTY" (si el Load Type es vacío).'),
    ('Acciones',        'Botón de tres puntos (⋮) que abre el menú de opciones para cada contenedor.'),
]
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
for i, h in enumerate(['Columna', 'Descripción']):
    tbl3.rows[0].cells[i].text = h
    for run in tbl3.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = BLANCO
    set_cell_bg(tbl3.rows[0].cells[i], '0A1B5B')
for col, desc in cols:
    row = tbl3.add_row().cells
    row[0].text = col; row[1].text = desc
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)

add_body('')
add_note('El Supervisor ve la tabla en modo solo lectura: no aparece el botón de opciones (⋮) '
         'ni el botón "Agregar nuevo contenedor".', '0A1B5B')

# ── 6.2 ─────────────────────────────────────────────────────────────────
add_h2('6.2  Agregar un nuevo contenedor — Paso 1: Información Completa')
add_body('Para registrar un nuevo contenedor, pulsar el botón azul "Agregar nuevo contenedor" '
         'ubicado en la parte superior derecha de la tabla. Se abrirá un formulario modal de tres pasos.')

add_h3('Barra de progreso')
add_body('En la parte superior del formulario se muestra una barra de progreso con los tres pasos:')
add_bullet('Paso 1 — Información Básica')
add_bullet('Paso 2 — Inspección de Trailer')
add_bullet('Paso 3 — Documentos')

add_h3('Campos del Paso 1')
add_body('El Paso 1 contiene dos secciones: Información Básica e Información Adicional.')
add_body('Campos obligatorios (marcados con *):', bold=True)

obligatorios = [
    ('TRAILER NO. *',   'Número único del tráiler. Campo de texto libre.'),
    ('TRAILER TYPE *',  'Tipo de tráiler. Se selecciona de una lista configurable por el Administrador.'),
    ('ACTUAL DATE *',   'Fecha real de llegada. Se selecciona con el selector de fecha.'),
    ('PO.# *',          'Número de orden de compra (Purchase Order). Se selecciona de lista configurable.'),
]
tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
for i, h in enumerate(['Campo', 'Descripción']):
    tbl4.rows[0].cells[i].text = h
    for run in tbl4.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(10); run.font.color.rgb=BLANCO
    set_cell_bg(tbl4.rows[0].cells[i], '0A1B5B')
for campo, desc in obligatorios:
    row = tbl4.add_row().cells
    row[0].text = campo; row[1].text = desc
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(10)

add_body('')
add_body('Campos opcionales de la sección Información Básica:', bold=True)
opcionales_basica = [
    ('SEA CONTAINER TYPE', 'Tipo de contenedor marítimo. Texto libre.'),
    ('USO EMBARQUES',      'Uso del embarque. Lista configurable.'),
    ('PORT OF ENTRY',      'Puerto de entrada. Lista configurable.'),
    ('LOAD TYPE',          'Tipo de carga (puede ser EMPTY). Lista configurable.'),
    ('STATUS',             'Estado del contenedor. Lista configurable.'),
    ('YARD / DESTINATION', 'Patio o destino del contenedor. Lista configurable.'),
    ('COMMENTS',           'Comentarios adicionales. Campo de texto largo.'),
    ('ATTACHMENTS',        'Adjuntar archivos desde esta sección (alternativa al Paso 3).'),
]
tbl5 = doc.add_table(rows=1, cols=2)
tbl5.style = 'Table Grid'
for i, h in enumerate(['Campo', 'Descripción']):
    tbl5.rows[0].cells[i].text = h
    for run in tbl5.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(10); run.font.color.rgb=BLANCO
    set_cell_bg(tbl5.rows[0].cells[i], '0A1B5B')
for campo, desc in opcionales_basica:
    row = tbl5.add_row().cells
    row[0].text = campo; row[1].text = desc
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(10)

add_body('')
add_body('Campos de la sección Información Adicional:', bold=True)
adicional = [
    ('QTY OF PALLETS',    'Cantidad total de pallets esperados.'),
    ('EMPTY DATE',        'Fecha de vaciado del contenedor.'),
    ('SEAL # (San Luis)', 'Número de sello en San Luis.'),
    ('DEPARTURE DATE',    'Fecha de salida.'),
    ('SEAL # (Yuma)',     'Número de sello en Yuma.'),
    ('AGING A',           'Valor de antigüedad A.'),
    ('ITEM TYPE',         'Tipo de artículo.'),
    ('AGING',             'Valor de antigüedad general.'),
    ('BOOKING #',         'Número de reservación o booking.'),
    ('DATE EXIT OF PORT', 'Fecha de salida del puerto.'),
]
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.style = 'Table Grid'
for i, h in enumerate(['Campo', 'Descripción']):
    tbl6.rows[0].cells[i].text = h
    for run in tbl6.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(10); run.font.color.rgb=BLANCO
    set_cell_bg(tbl6.rows[0].cells[i], '0A1B5B')
for campo, desc in adicional:
    row = tbl6.add_row().cells
    row[0].text = campo; row[1].text = desc
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(10)

add_body('')
add_h3('Botones del Paso 1')
add_bullet('"Guardar" — Guarda el Paso 1 sin avanzar. Útil para guardar el avance parcial.')
add_bullet('"Siguiente" — Valida los campos obligatorios y guarda antes de avanzar al Paso 2.')
add_note('Si falta un campo obligatorio, se muestra un mensaje de error y el sistema no avanza.')

# ── 6.3 ─────────────────────────────────────────────────────────────────
add_h2('6.3  Agregar un nuevo contenedor — Paso 2: Inspección de Trailer/Contenedor')
add_body('El Paso 2 registra la inspección física del tráiler al momento de su llegada. '
         'Contiene varias secciones.')

add_h3('Sección: Caja de Trailer')
campos_p2a = [
    ('CAJA DE TRAILER / TCN *', 'Número de caja o TCN. Campo obligatorio, texto libre.'),
    ('PLACAS',       'Número de placas del vehículo.'),
    ('ESTADO',       'Condición general (Bueno, Dañado, etc.).'),
    ('FECHA LLEGADA','Fecha de llegada física del tráiler. Selector de fecha.'),
]
for c, d in campos_p2a:
    add_bullet(d, bold_prefix=f'{c}: ')

add_h3('Sección: Turno')
add_body('Selección de turno mediante botones de opción (radio buttons):')
add_bullet('1er turno (6:45 AM – 6:45 PM)')
add_bullet('2do turno (6:45 PM – 6:45 AM)')

add_h3('Sección: Checklist de inspección')
add_body('Antes de iniciar la carga/descarga, el operador debe marcar los siguientes puntos '
         'de la lista de verificación:')
condiciones = [
    'Condiciones de las dos puertas del trailer',
    'Revisar que se encuentre libre de olores extraños',
    'Revisar que no tenga plagas, basura y humedad',
    'Revisar que los empaques de la carga están cerrados',
    'Condiciones de la pared del fondo del trailer',
    'Condiciones de paredes internas del trailer',
    'Condiciones internas del techo del trailer',
    'Condiciones de piso / plataforma interna del trailer',
]
for i, c in enumerate(condiciones, 1):
    add_bullet(c, level=0, bold_prefix=f'{i}. ')

add_h3('Sección: Información de Llegada')
campos_llegada = [
    ('SELLOS',       'Números de sellos de seguridad del tráiler.'),
    ('RAMPA',        'Número de rampa donde se descarga.'),
    ('HORA REGISTRO','Hora exacta del registro de llegada.'),
    ('TOTAL PALLETS','Cantidad de pallets recibidos físicamente.'),
]
for c, d in campos_llegada:
    add_bullet(d, bold_prefix=f'{c}: ')

add_h3('Sección: Longitud del Contenedor')
add_body('Se selecciona la longitud mediante botones de opción (radio buttons) de las opciones '
         'configuradas por el Administrador. Si se selecciona la opción "Otro", aparece un campo '
         'adicional para ingresar la medida exacta en pulgadas.')

add_h3('Sección: Origen (Arribo)')
add_body('Se selecciona el lugar de origen del contenedor mediante botones de opción, '
         'de las opciones configuradas en el catálogo "Origen".')

add_h3('Sección: Empresas')
add_body('Se pueden seleccionar una o varias empresas relacionadas con el contenedor mediante '
         'casillas de verificación (checkboxes). Las opciones provienen del catálogo "Empresas".')

add_h3('Sección: Firma del Responsable de Descarga')
add_body('Esta es una de las secciones más importantes del Paso 2:')
add_numbered('Seleccionar el responsable en el campo RESPONSIBLE (lista desplegable del catálogo).')
add_numbered('Firmar directamente en el área de firma con el dedo (en tabletas) o con el mouse. '
             'El área de firma es un lienzo interactivo de color blanco.')
add_numbered('Presionar "Guardar Firma" para confirmar la firma. Aparecerá el mensaje '
             '"✓ Firma guardada correctamente".')
add_numbered('Si se desea borrar y volver a firmar, presionar "Limpiar Firma".')

add_note('La firma se guarda como imagen digital (PNG) junto al registro del contenedor y '
         'queda incluida en el PDF generado.', '0A1B5B')

add_h3('Botones del Paso 2')
add_bullet('"Guardar" — Guarda el Paso 2 sin avanzar.')
add_bullet('"Atrás" — Regresa al Paso 1 sin perder los datos ingresados.')
add_bullet('"Siguiente" — Guarda el Paso 2 y avanza al Paso 3.')

# ── 6.4 ─────────────────────────────────────────────────────────────────
add_h2('6.4  Agregar un nuevo contenedor — Paso 3: Documentos')
add_body('El Paso 3 permite adjuntar fotografías y documentos al registro del contenedor.')

add_h3('Opciones disponibles')
add_bullet('"Capturar Foto" — Abre la cámara del dispositivo (disponible en tabletas y móviles '
           'con cámara). Permite tomar una foto directamente.')
add_bullet('"Subir Archivo" — Permite seleccionar uno o varios archivos desde el almacenamiento '
           'del dispositivo (formatos aceptados: imágenes y PDF).')

add_h3('Documentos guardados')
add_body('Si el contenedor ya tiene documentos cargados previamente, se muestran en la lista '
         '"Documentos guardados" con su nombre y un botón para eliminarlos individualmente (✕).')

add_h3('Nuevos documentos')
add_body('Los archivos seleccionados en la sesión actual aparecen en la lista "Nuevos documentos" '
         'con su nombre y tamaño en KB. Pueden removerse antes de guardar con el botón ✕.')

add_h3('Botones del Paso 3')
add_bullet('"Guardar" — Sube los documentos al servidor y cierra el formulario. Se requiere al '
           'menos un documento para completar el Paso 3.')
add_bullet('"Atrás" — Regresa al Paso 2 sin perder los archivos seleccionados.')

add_note('Después de guardar el Paso 3, el sistema muestra el mensaje: '
         '"¡Documentos guardados! Usa el botón Archivar para enviar al archivo." '
         'El contenedor quedará como "Completado" y listo para archivarse.', '10B981')

# ── 6.5 ─────────────────────────────────────────────────────────────────
add_h2('6.5  Editar un contenedor existente')
add_body('Para editar un contenedor que ya fue registrado:')
add_numbered('En la tabla de Contenedores, hacer clic sobre la fila del contenedor deseado.')
add_numbered('Se abrirá el formulario modal con los datos ya cargados en los campos correspondientes.')
add_numbered('Si el Paso 2 ya existe, el formulario abrirá directamente en el Paso 2.')
add_numbered('Modificar los campos necesarios.')
add_numbered('Presionar "Guardar" en el paso que se esté editando.')

add_note('El Supervisor no puede hacer clic en las filas para editar — solo puede ver la tabla.')

# ── 6.6 ─────────────────────────────────────────────────────────────────
add_h2('6.6  Archivar un contenedor')
add_body('Archivar mueve el contenedor del flujo activo al Archivo, indicando que el proceso '
         'operativo de ese contenedor ha concluido.')
add_numbered('En la tabla, localizar el contenedor a archivar.')
add_numbered('Hacer clic en el botón de tres puntos (⋮) de ese contenedor.')
add_numbered('En el menú desplegable, seleccionar "Archivar".')
add_numbered('Confirmar la acción en el diálogo: "¿Estás seguro de que deseas archivar este contenedor?"')
add_numbered('El contenedor desaparece de la lista activa y pasa al módulo Archivo.')

add_note('Solo Administradores y Operadores pueden archivar contenedores.')

# ── 6.7 ─────────────────────────────────────────────────────────────────
add_h2('6.7  Vaciar campos opcionales')
add_body('La opción "Vaciar" está disponible únicamente para contenedores que ya hayan sido '
         'archivados (se muestra en el menú ⋮ cuando el contenedor tiene el indicador "Archivado").')
add_body('Al vaciar, se borran todos los campos opcionales del contenedor, conservando únicamente:')
add_bullet('Trailer No.')
add_bullet('Trailer Type')
add_bullet('Actual Date')
add_bullet('PO.#')

add_body('Pasos:')
add_numbered('Abrir el menú ⋮ del contenedor archivado.')
add_numbered('Seleccionar "Vaciar".')
add_numbered('Confirmar en el diálogo de verificación.')

# ── 6.8 ─────────────────────────────────────────────────────────────────
add_h2('6.8  Eliminar un contenedor')
add_body('La eliminación borra el contenedor de forma permanente del sistema.')
add_numbered('Abrir el menú ⋮ del contenedor en la tabla.')
add_numbered('Seleccionar "Eliminar".')
add_numbered('Confirmar en el diálogo: "¿Estás seguro? Esto eliminará el contenedor de forma permanente".')

add_note('⚠  Esta acción no puede deshacerse. Verificar que sea el contenedor correcto '
         'antes de confirmar.', 'DC2626')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7 — ARCHIVO
# ═══════════════════════════════════════════════════════════════════════════
add_h1('7. Módulo: Archivo')
add_body('El módulo Archivo almacena todos los contenedores que han sido archivados (completados '
         'y enviados al archivo). Permite consultar, filtrar, visualizar en PDF, editar y eliminar registros.')

add_body('Las columnas de la tabla del Archivo son:')
cols_arch = [
    ('FECHA',           'Hora del día en que se creó el registro (agrupados por fecha).'),
    ('TRAILER NO.',     'Número del tráiler.'),
    ('CONTENEDOR',      'Sea Container Type.'),
    ('PUERTO DE ENTRADA','Puerto de entrada declarado.'),
    ('LOAD TYPE',       'Tipo de carga.'),
    ('RESPONSABLE',     'Responsable de descarga registrado en el Paso 2.'),
    ('ESTADO',          '"Completo" (tiene documentos adjuntos) o "Sin documentos".'),
    ('ACCIONES',        'Botones: Ver PDF, Editar y Eliminar.'),
]
tbl7 = doc.add_table(rows=1, cols=2)
tbl7.style = 'Table Grid'
for i, h in enumerate(['Columna', 'Descripción']):
    tbl7.rows[0].cells[i].text = h
    for run in tbl7.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(10); run.font.color.rgb=BLANCO
    set_cell_bg(tbl7.rows[0].cells[i], '0A1B5B')
for c, d in cols_arch:
    row = tbl7.add_row().cells
    row[0].text = c; row[1].text = d
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(10)

add_body('')

# ── 7.1 ─────────────────────────────────────────────────────────────────
add_h2('7.1  Filtros disponibles')
add_body('En la parte superior del Archivo se encuentran los siguientes filtros:')
filtros = [
    ('Rango de fechas (Desde / Hasta)',
     'Filtra por la fecha de creación del registro. Por defecto muestra el día de hoy. '
     'El botón de actualizar (↺) restablece las fechas a hoy.'),
    ('Tipo',
     'Filtra por el tipo de embarque: Todos los tipos / Inspección / Llegada / Otros.'),
    ('Estado',
     'Filtra por si tiene documentos: Todos / Completo / Incompleto.'),
    ('Buscador',
     'Búsqueda por texto en Trailer No., Sea Container Type o Puerto de Entrada.'),
]
for nombre, desc in filtros:
    add_bullet(desc, bold_prefix=f'{nombre}: ')

# ── 7.2 ─────────────────────────────────────────────────────────────────
add_h2('7.2  Generar y descargar PDF')
add_body('Para ver el reporte completo de un contenedor archivado en formato PDF:')
add_numbered('Localizar el contenedor en la tabla del Archivo.')
add_numbered('Hacer clic en el ícono de PDF (🖹) en la columna Acciones.')
add_numbered('El sistema genera automáticamente un PDF con todos los datos del contenedor:')
add_bullet('Información del Paso 1 (datos básicos e información adicional).', level=1)
add_bullet('Resultados de la inspección del Paso 2 (checklist con condiciones marcadas).', level=1)
add_bullet('Firma del responsable de descarga (imagen de la firma capturada).', level=1)
add_bullet('Lista de documentos adjuntos del Paso 3.', level=1)
add_numbered('El PDF se descarga automáticamente en el dispositivo.')

# ── 7.3 ─────────────────────────────────────────────────────────────────
add_h2('7.3  Editar registro archivado')
add_body('Para modificar los datos de un contenedor que ya fue archivado:')
add_numbered('Hacer clic en el ícono de edición (✏) en la fila del contenedor.')
add_numbered('Se abrirá el formulario modal con todos los datos cargados.')
add_numbered('Modificar los campos necesarios y guardar.')

add_note('Esta opción no está disponible para Supervisores.', '0A1B5B')

# ── 7.4 ─────────────────────────────────────────────────────────────────
add_h2('7.4  Eliminar del archivo')
add_body('Para eliminar un registro del archivo:')
add_numbered('Hacer clic en el ícono de basura (🗑) en la fila del contenedor.')
add_numbered('Confirmar en el diálogo: "¿Eliminar este registro del archivo?".')
add_numbered('El registro se elimina permanentemente.')

add_note('⚠  La eliminación en el Archivo es permanente. Solo disponible para '
         'Administradores y Operadores.', 'DC2626')

# ── 7.5 ─────────────────────────────────────────────────────────────────
add_h2('7.5  Exportar a Excel')
add_body('En la parte superior del Archivo existe el botón "Exportar a Excel" que descarga '
         'todos los contenedores actualmente visibles (según los filtros aplicados) en un '
         'archivo .xlsx listo para análisis en Microsoft Excel u otras herramientas de hoja de cálculo.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 8 — REPORTES
# ═══════════════════════════════════════════════════════════════════════════
add_h1('8. Módulo: Reportes')
add_body('El módulo de Reportes presenta un dashboard visual con estadísticas e indicadores '
         'clave del flujo operativo de contenedores. Es accesible para todos los roles.')

add_h3('Selector de período')
add_body('En la parte superior del dashboard hay tres botones para filtrar el período de tiempo:')
add_bullet('"Hoy" — Muestra datos del día actual.')
add_bullet('"Esta semana" — Muestra datos de la semana en curso.')
add_bullet('"Este mes" — Muestra datos del mes actual.')

# ── 8.1 ─────────────────────────────────────────────────────────────────
add_h2('8.1  KPIs — Tarjetas de indicadores')
add_body('La fila superior muestra cuatro tarjetas KPI (Key Performance Indicators) con '
         'los datos del período seleccionado:')
kpis = [
    ('Total Contenedores (azul)',   'Número total de contenedores registrados en el período.'),
    ('Completados (verde)',         'Número de contenedores con status "Completado", más el porcentaje que representan.'),
    ('En Proceso (amarillo)',       'Número de contenedores aún en proceso.'),
    ('Pallets Totales (morado)',    'Suma de todos los pallets registrados en el campo Total Pallets.'),
]
for k, d in kpis:
    add_bullet(d, bold_prefix=f'{k}: ')

# ── 8.2 ─────────────────────────────────────────────────────────────────
add_h2('8.2  Panel expandible de detalle')
add_body('Al hacer clic en cualquiera de las primeras tres tarjetas KPI, se despliega un '
         'panel de detalle debajo de la fila de KPIs. Este panel muestra la lista de '
         'contenedores que corresponden al KPI seleccionado, con columnas:')
add_bullet('Trailer — Número de tráiler.')
add_bullet('Origen — Lugar de origen.')
add_bullet('Turno — 1er o 2do turno.')
add_bullet('Pallets — Cantidad de pallets.')
add_bullet('Estado — Badge "✓ Completado" o "● En proceso".')
add_body('Hacer clic de nuevo en el mismo KPI o en el botón ✕ del panel cierra el detalle.')
add_body('Al hacer clic en la tarjeta "Pallets Totales", el panel muestra:')
add_bullet('Badges con el total de pallets por origen.')
add_bullet('Barras horizontales por origen con la proporción de pallets.')
add_bullet('Lista detallada de contenedores con pallets ordenados de mayor a menor.')

# ── 8.3 ─────────────────────────────────────────────────────────────────
add_h2('8.3  Progreso de completado (gráfica de anillo)')
add_body('Una gráfica circular (anillo) muestra visualmente el porcentaje de contenedores '
         'completados vs. en proceso en el período seleccionado. El número en el centro '
         'indica el porcentaje de completado.')

# ── 8.4 ─────────────────────────────────────────────────────────────────
add_h2('8.4  Contenedores por mes')
add_body('Un gráfico de barras verticales muestra la cantidad de contenedores registrados '
         'en los últimos seis meses, permitiendo identificar tendencias operativas.')

# ── 8.5 ─────────────────────────────────────────────────────────────────
add_h2('8.5  Gráficas de barras (Turno / Origen / Empresa)')
add_body('Tres gráficas de barras horizontales muestran la distribución de contenedores por:')
add_bullet('"Por Turno" — Comparación entre 1er turno y 2do turno (barras azules).')
add_bullet('"Por Origen" — Top 6 orígenes con mayor número de contenedores (barras verdes).')
add_bullet('"Por Empresa" — Top 6 empresas con mayor número de contenedores (barras amarillas).')
add_body('El largo de cada barra es proporcional al valor más alto. El número al final de cada '
         'barra indica el total exacto.')

add_h3('Botón Actualizar')
add_body('El botón "Actualizar" en la parte superior del dashboard recarga los datos del servidor '
         'sin cambiar el período seleccionado.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 9 — ENTREGA DE TURNO
# ═══════════════════════════════════════════════════════════════════════════
add_h1('9. Módulo: Entrega de Turno')
add_body('El módulo Entrega de Turno muestra un resumen de todos los contenedores completados '
         'y archivados durante un turno específico. Facilita el proceso de cambio de turno '
         'entre operadores y supervisores.')

# ── 9.1 ─────────────────────────────────────────────────────────────────
add_h2('9.1  Visualizar el turno activo')
add_body('Al abrir el módulo, el sistema detecta automáticamente el turno actual basado '
         'en la hora del servidor:')
add_bullet('1er turno: 6:45 AM – 6:45 PM')
add_bullet('2do turno: 6:45 PM – 6:45 AM')

add_body('Los filtros disponibles son:')
add_bullet('"Fecha" — Selector de fecha. Por defecto muestra el día de hoy.', bold_prefix='')
add_bullet('"Turno" — Selector desplegable: 1er turno o 2do turno.')
add_bullet('"Total contenedores" — Badge que muestra en tiempo real la cantidad de contenedores del turno.')

add_body('La tabla principal muestra los contenedores completados del turno con columnas:')
cols_turno = [
    ('Rampa',              'Número de rampa donde se descargó el tráiler.'),
    ('Tráiler No.',        'Número identificador del tráiler.'),
    ('Caja Tráiler',       'Número de caja o TCN.'),
    ('Contenedor',         'Tipo de contenedor marítimo.'),
    ('Pallets recibidos',  'Total de pallets recibidos (campo TotalPallets del Paso 2).'),
    ('Capacidad (Pallets)','Capacidad declarada (campo QtyPallets del Paso 1).'),
    ('% Uso',              'Porcentaje calculado: (Pallets recibidos / Capacidad) × 100.'),
    ('Cliente / Empresa',  'Empresas asociadas al contenedor.'),
    ('Comentario',         'Comentarios del Paso 1.'),
]
tbl8 = doc.add_table(rows=1, cols=2)
tbl8.style = 'Table Grid'
for i, h in enumerate(['Columna', 'Descripción']):
    tbl8.rows[0].cells[i].text = h
    for run in tbl8.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(9); run.font.color.rgb=BLANCO
    set_cell_bg(tbl8.rows[0].cells[i], '0A1B5B')
for c, d in cols_turno:
    row = tbl8.add_row().cells
    row[0].text = c; row[1].text = d
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(9)

add_body('')
add_body('Al pie de la tabla aparece el total de tráilas descargadas en el turno.')

# ── 9.2 ─────────────────────────────────────────────────────────────────
add_h2('9.2  Guardar entrega de turno')
add_body('Para registrar formalmente una entrega de turno y conservar el historial:')
add_numbered('Verificar que la fecha y el turno sean los correctos en los filtros.')
add_numbered('Confirmar que la tabla muestra los contenedores esperados.')
add_numbered('Hacer clic en el botón azul "Guardar entrega".')
add_numbered('El sistema muestra el mensaje de confirmación con el total de tráilas guardadas.')
add_numbered('La entrega queda registrada en la sección "Entregas de turno guardadas".')

add_note('Si no hay contenedores en el turno seleccionado, el sistema mostrará el aviso '
         '"No hay contenedores para guardar en este turno".', 'F59E0B')
add_note('Esta acción no está disponible para el rol Supervisor.', '0A1B5B')

# ── 9.3 ─────────────────────────────────────────────────────────────────
add_h2('9.3  Exportar entrega a Excel')
add_body('Las entregas guardadas se muestran en la sección inferior "Entregas de turno guardadas" '
         'con las columnas: Fecha guardado, Fecha turno, Turno, Total tráilas y Acciones.')
add_body('Para exportar una entrega guardada a Excel:')
add_numbered('Localizar la entrega en la tabla de entregas guardadas.')
add_numbered('Hacer clic en el botón "Exportar a Excel" (ícono de descarga).')
add_numbered('Se descarga un archivo .xlsx con el detalle completo de los contenedores '
             'de ese turno.')

# ── 9.4 ─────────────────────────────────────────────────────────────────
add_h2('9.4  Eliminar entrega guardada')
add_body('Para eliminar una entrega del historial:')
add_numbered('Localizar la entrega en la tabla de entregas guardadas.')
add_numbered('Hacer clic en el ícono de basura (🗑).')
add_numbered('Confirmar en el diálogo: "¿Eliminar esta entrega guardada?".')
add_numbered('La entrega se elimina del historial.')

add_note('Esta acción no está disponible para el rol Supervisor.', '0A1B5B')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 10 — ADMINISTRACIÓN
# ═══════════════════════════════════════════════════════════════════════════
add_h1('10. Panel de Administración')
add_body('El Panel de Administración es accesible únicamente para usuarios con rol Administrador. '
         'Contiene dos secciones principales: Configuración de Listas y Gestión de Usuarios.')
add_note('Si un usuario sin rol de Administrador intenta acceder a esta sección, verá el mensaje '
         '"Acceso restringido — Solo usuarios con rol Administrador pueden acceder a esta sección."', 'DC2626')

# ── 10.1 ────────────────────────────────────────────────────────────────
add_h2('10.1  Configuración de Listas')
add_body('Las listas configurables son los catálogos de opciones que aparecen en los '
         'selectores del formulario de contenedores. El Administrador puede agregar o '
         'eliminar opciones de cada lista.')

add_body('Listas disponibles:')
listas = [
    ('Trailer Type',         'Tipos de tráiler disponibles para seleccionar en el Paso 1.'),
    ('Uso Embarques',        'Usos del embarque (Llegada, Inspección, Otros, etc.).'),
    ('Port of Entry',        'Puertos de entrada disponibles.'),
    ('Load Type',            'Tipos de carga (FULL, EMPTY, PARTIAL, etc.).'),
    ('Status',               'Estados del contenedor.'),
    ('Yard / Destination',   'Patios o destinos disponibles.'),
    ('Responsible',          'Responsables de descarga disponibles para seleccionar en la firma.'),
    ('P.O. #',               'Números de Purchase Order disponibles.'),
    ('Empresas',             'Empresas que pueden asociarse a un contenedor.'),
    ('Origen (Arribo)',      'Lugares de origen disponibles para seleccionar en el Paso 2.'),
    ('Longitud Contenedor',  'Longitudes estándar del contenedor (40 ft, 45 ft, 53 ft, Otro, etc.).'),
]
tbl9 = doc.add_table(rows=1, cols=2)
tbl9.style = 'Table Grid'
for i, h in enumerate(['Lista', 'Descripción']):
    tbl9.rows[0].cells[i].text = h
    for run in tbl9.rows[0].cells[i].paragraphs[0].runs:
        run.bold=True; run.font.size=Pt(10); run.font.color.rgb=BLANCO
    set_cell_bg(tbl9.rows[0].cells[i], '0A1B5B')
for l, d in listas:
    row = tbl9.add_row().cells
    row[0].text = l; row[1].text = d
    for cell in row:
        for run in cell.paragraphs[0].runs: run.font.size = Pt(10)

add_body('')
add_h3('Agregar una nueva opción a una lista')
add_numbered('Ir a Administración → Configuración de Listas.')
add_numbered('Seleccionar la lista deseada en la fila de pestañas (Trailer Type, Uso Embarques, etc.).')
add_numbered('En el campo "Nombre de la opción", escribir el texto de la nueva opción.')
add_numbered('Presionar el botón "Agregar" o la tecla Enter.')
add_numbered('La opción aparece inmediatamente en la lista de "Opciones actuales".')

add_h3('Eliminar una opción de una lista')
add_numbered('Localizar la opción en la lista de "Opciones actuales".')
add_numbered('Hacer clic en el ícono de basura (🗑) junto a la opción.')
add_numbered('Confirmar en el diálogo de verificación.')
add_note('⚠  Eliminar una opción de una lista no afecta los registros ya guardados, pero '
         'esa opción ya no estará disponible para nuevos registros.', 'F59E0B')

add_h3('Actualizar la lista')
add_body('El botón de actualizar (↺) en el encabezado de opciones recarga las opciones '
         'desde el servidor para reflejar cambios recientes.')

# ── 10.2 ────────────────────────────────────────────────────────────────
add_h2('10.2  Gestión de Usuarios')
add_body('La sección de Usuarios permite al Administrador crear nuevas cuentas y activar '
         'o desactivar usuarios existentes.')

add_h3('Crear un nuevo usuario')
add_numbered('Ir a Administración → Usuarios.')
add_numbered('Completar el formulario "Agregar nuevo usuario":')
add_bullet('Nombre completo — Nombre real del usuario.', level=1)
add_bullet('Email — Correo electrónico que servirá como nombre de usuario para iniciar sesión.', level=1)
add_bullet('Contraseña — Contraseña inicial. Se puede mostrar/ocultar con el ícono de ojo (👁).', level=1)
add_bullet('Rol — Seleccionar entre Admin, Supervisor u Operador.', level=1)
add_numbered('Al seleccionar un rol, aparece una descripción del rol con sus permisos.')
add_numbered('Hacer clic en "Crear usuario".')
add_numbered('Si todos los campos son válidos, el usuario se crea y aparece en la lista inferior.')

add_h3('Lista de usuarios registrados')
add_body('En la parte inferior de la sección se muestra la lista de todos los usuarios '
         'del sistema con:')
add_bullet('Nombre completo.')
add_bullet('Correo electrónico.')
add_bullet('Etiqueta de rol (Admin / Supervisor / Operador).')
add_bullet('Estado: "Activo" o "Inactivo".')
add_bullet('Botón para activar/desactivar.')

add_h3('Activar o desactivar un usuario')
add_body('Para controlar el acceso de un usuario sin eliminarlo del sistema:')
add_numbered('Localizar al usuario en la lista.')
add_numbered('Hacer clic en el ícono de persona junto a su nombre:')
add_bullet('Ícono "Desactivar" (persona con X) — desactiva al usuario activo.', level=1)
add_bullet('Ícono "Activar" (persona normal) — reactiva a un usuario inactivo.', level=1)
add_numbered('Confirmar en el diálogo de verificación.')
add_numbered('El estado del usuario se actualiza inmediatamente en la lista.')

add_note('Un usuario desactivado no puede iniciar sesión aunque conozca sus credenciales. '
         'Sus registros históricos se conservan.', '0A1B5B')

add_h3('Descripción de roles en el formulario')
add_body('Al seleccionar un rol en el formulario de nuevo usuario, se muestran tarjetas '
         'informativas con el color y descripción de los tres roles disponibles, '
         'facilitando la decisión de qué rol asignar.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECCIÓN 11 — PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════════════════════════
add_h1('11. Preguntas frecuentes')

faqs = [
    ('¿Qué hago si olvidé mi contraseña?',
     'Contactar al Administrador del sistema para que restablezca la contraseña. '
     'El Administrador puede crear una nueva contraseña al editar o recrear el usuario.'),

    ('¿Puedo registrar un contenedor sin completar todos los pasos en un solo momento?',
     'Sí. Al presionar "Guardar" en cualquier paso, el progreso se guarda. El contenedor '
     'aparece en la tabla activa y puede abrirse más tarde para continuar con el Paso 2 o el Paso 3.'),

    ('¿La firma se puede modificar después de guardar?',
     'Sí. Al editar el contenedor desde la tabla o desde el Archivo, se puede ir al Paso 2, '
     'limpiar la firma y trazar una nueva, luego guardar.'),

    ('¿Qué pasa si subo el documento incorrecto en el Paso 3?',
     'Antes de guardar, se puede eliminar el archivo con el botón ✕ en la lista de '
     '"Nuevos documentos". Si ya fue guardado, abrir el contenedor para editar, ir al Paso 3 '
     'y eliminar el documento de "Documentos guardados" con el botón ✕.'),

    ('¿Por qué no veo el botón "Agregar nuevo contenedor"?',
     'Ese botón solo es visible para roles Administrador y Operador. Si tienes rol Supervisor, '
     'el módulo es solo de consulta.'),

    ('¿Cómo sé que un contenedor ya fue completado y archivado?',
     'En la tabla de Contenedores, el status mostrará "Completado". Si ya fue archivado, '
     'aparecerá un indicador "Archivado" bajo el status y el contenedor también estará '
     'visible en el módulo Archivo.'),

    ('¿Por qué el Load Type aparece como "EMPTY" en la tabla?',
     'Cuando el campo Load Type está configurado como EMPTY, el sistema muestra un '
     'badge especial "EMPTY" en lugar del status normal, indicando que el contenedor '
     'fue recibido vacío.'),

    ('¿Puedo exportar los reportes a Excel?',
     'El módulo Reportes no tiene exportación directa, pero el módulo Archivo tiene '
     'el botón "Exportar a Excel" que descarga todos los registros visibles. '
     'El módulo Entrega de turno también permite exportar cada entrega guardada a Excel.'),

    ('¿Qué son las listas configurables?',
     'Son los catálogos de opciones que aparecen en los selectores del formulario de '
     'contenedores (Trailer Type, Origen, Empresas, etc.). Solo el Administrador puede '
     'agregar o eliminar opciones desde el Panel de Administración.'),

    ('¿El sistema funciona en dispositivos móviles?',
     'Sí. La interfaz es responsiva y se adapta a tabletas y móviles. '
     'La captura de firma táctil está optimizada para pantallas táctiles. '
     'En pantallas pequeñas, la barra lateral se oculta y se accede con el menú ☰.'),
]

for pregunta, respuesta in faqs:
    add_h3(f'• {pregunta}')
    add_body(respuesta)
    add_body('')

doc.add_page_break()

# ── PIE DE PÁGINA ────────────────────────────────────────────────────────
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_pie = p_pie.add_run(
    'Flextronics Technologies  ·  San Luis Río Colorado, Sonora, México\n'
    'Sistema de Gestión de Recibos  ·  v1.0  ·  2026\n'
    'Documento de uso interno — Todos los derechos reservados')
r_pie.font.size = Pt(9)
r_pie.font.color.rgb = RGBColor(0x77,0x77,0x77)

# ── GUARDAR ──────────────────────────────────────────────────────────────
output_path = r'C:\Users\elerv\Desktop\Manual_FlexWebApp.docx'
doc.save(output_path)
print(f'Manual guardado en: {output_path}')
